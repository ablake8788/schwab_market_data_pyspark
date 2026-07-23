"""
reporting/report_builder.py
-----------------------------
Builds the Bollinger Bands / Z-Score summary as a .docx using python-docx.
Pure presentation logic — takes plain dicts/lists from reporting/queries.py,
returns nothing but a saved file. No SQL, no SMTP.
"""

from __future__ import annotations

from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


def _set_cell_shading(cell, color_hex: str) -> None:
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    cell._tc.get_or_add_tcPr().append(shd)


def _add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    for j, header in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.width = Inches(widths[j])
        run = cell.paragraphs[0].add_run(header)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _set_cell_shading(cell, "374151")

    for i, row in enumerate(rows, start=1):
        for j, value in enumerate(row):
            cell = table.rows[i].cells[j]
            cell.width = Inches(widths[j])
            run = cell.paragraphs[0].add_run(str(value))
            run.font.size = Pt(9)


def _fmt_date(value) -> str:
    if value is None:
        return "-"
    return value.strftime("%Y-%m-%d %H:%M") if hasattr(value, "strftime") else str(value)


def _fmt_num(value, decimals: int = 2) -> str:
    if value is None:
        return "-"
    return f"{float(value):.{decimals}f}"


def build_summary_docx(
    overall: dict,
    bollinger_summary: list[dict],
    zscore_summary: list[dict],
    top_anomalies: list[dict],
    output_path: str | Path,
    anomaly_chart_path: str | Path | None = None,
    symbol_charts: list[tuple[str, str | Path]] | None = None,
) -> Path:
    """
    Builds the report and saves it to output_path. Returns the resolved Path.

    anomaly_chart_path : PNG from charts.build_anomaly_count_chart(), inserted
        after the Per-Symbol Summary table. Omitted entirely if None.
    symbol_charts : [(symbol, PNG path), ...] from charts.build_symbol_chart(),
        one per symbol, inserted in a "Symbol Charts" section at the end.
    """
    output_path = Path(output_path)

    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    title = doc.add_heading("Schwab Market Data — Bollinger Bands & Z-Score Summary", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph(
        f"Generated {datetime.now(timezone.utc):%Y-%m-%d %H:%M UTC} — "
        f"schwab_market_data_pyspark, dbo.SchwabQuotesHistory_SparkBollinger / _SparkZScore"
    )
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in sub.runs:
        r.font.size = Pt(9.5)
        r.italic = True
        r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()

    # ── Overall totals ──────────────────────────
    doc.add_heading("Overall", level=2)
    _add_table(
        doc,
        headers=["Metric", "Value"],
        rows=[
            ["Symbols", overall.get("SymbolCount", "-")],
            ["Date range", f"{_fmt_date(overall.get('MinDate'))} → {_fmt_date(overall.get('MaxDate'))}"],
            ["Bollinger rows (excl. warmup)", overall.get("TotalBollingerRows", "-")],
            ["Z-Score rows (excl. warmup)", overall.get("TotalZScoreRows", "-")],
            ["Anomalies flagged", overall.get("TotalAnomalies", "-")],
        ],
        widths=[3.5, 3.5],
    )

    doc.add_paragraph()

    # ── Per-symbol summary (merge Bollinger + Z-Score by Symbol) ──
    doc.add_heading("Per-Symbol Summary", level=2)
    zscore_by_symbol = {row["Symbol"]: row for row in zscore_summary}

    rows = []
    for b in bollinger_summary:
        z = zscore_by_symbol.get(b["Symbol"], {})
        rows.append([
            b["Symbol"],
            b["BarCount"],
            _fmt_num(b["MinClose"]),
            _fmt_num(b["MaxClose"]),
            _fmt_num(b["AvgClose"]),
            z.get("AnomalyCount", 0),
            _fmt_num(z.get("MaxAbsZScore"), 2),
        ])

    _add_table(
        doc,
        headers=["Symbol", "Bars", "Min Close", "Max Close", "Avg Close", "Anomalies", "Max |Z|"],
        rows=rows,
        widths=[0.8, 0.7, 0.9, 0.9, 0.9, 0.9, 0.8],
    )

    doc.add_paragraph()

    if anomaly_chart_path is not None:
        pic_p = doc.add_paragraph()
        pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pic_p.add_run().add_picture(str(anomaly_chart_path), width=Inches(6.5))
        doc.add_paragraph()

    # ── Top anomalies ───────────────────────────
    doc.add_heading("Most Extreme Anomalies (Top 15 by |Z-Score|)", level=2)
    if top_anomalies:
        _add_table(
            doc,
            headers=["Symbol", "Bar Date/Time", "Close Price", "Z-Score"],
            rows=[
                [a["Symbol"], _fmt_date(a["BarDateTime"]), _fmt_num(a["ClosePrice"]), _fmt_num(a["ZScore"], 3)]
                for a in top_anomalies
            ],
            widths=[1.0, 2.0, 1.5, 1.5],
        )
    else:
        doc.add_paragraph("No anomalies flagged in the current data.")

    # ── Per-symbol charts ────────────────────────
    if symbol_charts:
        doc.add_paragraph()
        doc.add_heading("Symbol Charts", level=2)
        for symbol, chart_path in symbol_charts:
            heading = doc.add_paragraph()
            heading_run = heading.add_run(symbol)
            heading_run.bold = True
            heading_run.font.size = Pt(11)

            pic_p = doc.add_paragraph()
            pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pic_p.add_run().add_picture(str(chart_path), width=Inches(6.5))

    footer = doc.add_paragraph()
    footer_run = footer.add_run("Generated by Claude Code.")
    footer_run.italic = True
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path.resolve()
