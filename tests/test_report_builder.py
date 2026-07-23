from datetime import datetime

from docx import Document

from reporting.report_builder import build_summary_docx


def test_build_summary_docx_creates_valid_file(tmp_path):
    overall = {
        "SymbolCount": 2,
        "MinDate": datetime(2025, 10, 24),
        "MaxDate": datetime(2025, 12, 11),
        "TotalBollingerRows": 100,
        "TotalZScoreRows": 100,
        "TotalAnomalies": 3,
    }
    bollinger_summary = [
        {"Symbol": "AAPL", "BarCount": 50, "MinClose": 150.0, "MaxClose": 160.0, "AvgClose": 155.0},
        {"Symbol": "LUNR", "BarCount": 50, "MinClose": 10.0, "MaxClose": 14.0, "AvgClose": 12.0},
    ]
    zscore_summary = [
        {"Symbol": "AAPL", "AnomalyCount": 1, "MaxAbsZScore": 2.1},
        {"Symbol": "LUNR", "AnomalyCount": 2, "MaxAbsZScore": 3.4},
    ]
    top_anomalies = [
        {"Symbol": "LUNR", "BarDateTime": datetime(2025, 11, 1, 9, 30), "ClosePrice": 13.5, "ZScore": 3.4},
    ]

    out_path = tmp_path / "report.docx"
    result = build_summary_docx(overall, bollinger_summary, zscore_summary, top_anomalies, out_path)

    assert result.exists()
    doc = Document(str(result))
    assert len(doc.tables) == 3  # overall, per-symbol, top anomalies
    assert doc.tables[1].rows[1].cells[0].text == "AAPL"


def test_build_summary_docx_handles_no_anomalies(tmp_path):
    overall = {"SymbolCount": 1, "MinDate": None, "MaxDate": None, "TotalBollingerRows": 0,
               "TotalZScoreRows": 0, "TotalAnomalies": 0}

    out_path = tmp_path / "report.docx"
    result = build_summary_docx(overall, [], [], [], out_path)

    assert result.exists()
    doc = Document(str(result))
    # "Overall" + a header-only "Per-Symbol Summary" table; the "Top
    # Anomalies" table is skipped entirely (falls back to a paragraph)
    # when top_anomalies is empty.
    assert len(doc.tables) == 2
